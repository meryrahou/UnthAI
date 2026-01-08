from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()

    # Style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('NLP Preprocessing Guide: Algerian TikTok Data', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Technical strategy for Algerian TikTok sentiment and intent analysis.')

    # Section 1
    doc.add_heading('1. Structural Cleaning', level=1)
    
    doc.add_heading('GIFs & Stickers', level=2)
    doc.add_paragraph('Remove tokens like [GIF] or [Sticker]. These carry no semantic content.')
    
    doc.add_heading('Tags & User IDs', level=2)
    doc.add_paragraph('If a comment consists solely of tagged users (e.g., @username), delete the row. If tags appear within a sentence, strip them to clean the text.')
    
    doc.add_heading('Reply Prefixes', level=2)
    doc.add_paragraph('Strip TikTok "Replying to" prefixes using regex: ^(@[\\w.]+[: ]*|Replying to @[\\w.]+[: ]*).')

    doc.add_heading('Final Dataset Preparation', level=2)
    doc.add_paragraph('To ensure the dataset is ready for model training:')
    doc.add_paragraph('Text Deduplication: Drop duplicate rows based on the comment_text column to remove redundant signals.', style='List Bullet')
    doc.add_paragraph('Shuffling: Randomly shuffle the final rows to ensure that different restaurant sentiments are mixed, avoiding bias during training batches.', style='List Bullet')
    doc.add_paragraph('ID Reset: After shuffling and deduplicating, assign a new sequential id column (e.g., 1 to N).', style='List Bullet')

    # Section 2
    doc.add_heading('2. Text Normalization', level=1)
    
    doc.add_heading('Punctuation Intensity', level=2)
    doc.add_paragraph('Convert repeated punctuation like !!!!!!! or ?????? into a single [INTENSE] token. Do not delete them, as they signify strong intent.')

    doc.add_heading('Emoji Selective Mapping', level=2)
    doc.add_paragraph('Map high-signal emojis to tokens and delete all noise (random objects, flags, animals).')

    # Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Signals'
    hdr_cells[2].text = 'Token'

    data = [
        ('Positive', '❤️ 😍 🔥 😋 🤤 👌 💯 ✨ 🌟 🔝 👍 👏 🥰 🥘 🍔 🍕 🥙 🥗 🌮 🍗 🍡 🍱 🥧 🍰 🍦 🥂', '[POS_EMOJI]'),
        ('Negative', '🤮 🤢 😡 👎 💸 📉 🚫 💀 💩 🤡 🙄 😤 🤬 🚮 💔', '[NEG_EMOJI]'),
        ('Neutral', '📍 📞 🕒 🚗 🛵 🍴 ☕ 🥤 🍨 🥖 🧂', '[NEUT_EMOJI]')
    ]

    for cat, sig, tok in data:
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = sig
        row_cells[2].text = tok

    # Section 3
    doc.add_heading('3. Language Script Filtering', level=1)
    doc.add_paragraph('Keep only Arabic and Latin (French/English/Arabizi) scripts. Delete comments containing Cyrillic, Asian, or other non-target scripts.')

    # Section 4
    doc.add_heading('4. Intent & Multimodal Classification', level=1)
    doc.add_paragraph('Identify the Intent of the user to understand the driver of the feedback.')
    
    doc.add_paragraph('Categories:', style='List Bullet')
    doc.add_paragraph('Review: Personal experience sharing.', style='List Bullet')
    doc.add_paragraph('Inquiry: Asking for info (price, location).', style='List Bullet')
    doc.add_paragraph('Recommendation: Promoting or warning against.', style='List Bullet')
    doc.add_paragraph('Complaint: Expressing a specific failure.', style='List Bullet')
    doc.add_paragraph('Appreciation: Generalized praise.', style='List Bullet')

    # Note
    p = doc.add_paragraph()
    run = p.add_run('Sentiment Fairness: For "unknown" or general topics, classify intent as "General Appreciation" to ensure sentiment data is captured for every usable comment.')
    run.bold = True
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)

    # Save
    doc.save('PREPROCESSING_GUIDE.docx')
    print("Document updated: PREPROCESSING_GUIDE.docx")

if __name__ == "__main__":
    create_doc()
